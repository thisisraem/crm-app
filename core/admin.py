# Your imports here.
from django.contrib import admin
from django.http import HttpResponse
import xlwt
import csv
from django.utils.translation import gettext_lazy as _
import io
from docx import Document
from .models import Settings, CustomerRecords, Malls, Brands
from django.utils import timezone

# Register your models here.
admin.site.register(Settings)
admin.site.register(Malls)
admin.site.register(Brands)

# Header settings
admin.site.site_header = "NOVCO CRM AZE"
admin.site.index_title = "Welcome, CRM_AZE_ADMIN"

# OperatoCodeFilter & CustomerRecordsAdmin here.
class OperatorCodeFilter(admin.SimpleListFilter):
    title = _('Operator Code')
    parameter_name = 'operator-code'

    def lookups(self, request, model_admin):
        return (
            ('50', _('050')),
            ('51', _('051')),
            ('55', _('055')),
            ('99', _('099')),
            ('70', _('070')),
            ('77', _('077')),
        )

    def queryset(self, request, queryset):
        value = self.value()
        if value:
            return queryset.filter(mobile_number__startswith=value)
        return queryset
    
class CustomerRecordsAdmin(admin.ModelAdmin):
    list_display = ('name', 'gender', 'age', 'mobile_number', 'malls_display', 'get_brands', 'display_comments', 'created_at_short', 'is_checked')
    list_filter = ('gender', 'age', 'malls', 'brands', OperatorCodeFilter, 'created_at', 'is_checked')
    search_fields = ('name', 'surname', 'mobile_number')
    list_per_page = 15

    def display_comments(self, obj):
        character_limit = 25
        comments = obj.comments
        
        if len(comments) > character_limit:
            comments = comments[:character_limit] + '...'
        
        return comments
    display_comments.short_description = 'Comments'

    def malls_display(self, obj):
        return ", ".join([str(mall) for mall in obj.malls.all()])
    malls_display.short_description = 'Malls'

    def get_brands(self, obj):
        return ", ".join([brand.name for brand in obj.brands.all()])

    get_brands.short_description = 'Brands'

    def created_at_short(self, obj):
        created_at = timezone.localtime(obj.created_at)
        return created_at.strftime('%m/%d/%y - %H:%M')

    created_at_short.short_description = 'Created At'
  
    actions = [
        'export_selected_customers_xls',
        'export_selected_customers_xls_name_phone',
        'export_selected_customers_csv',
        'export_selected_customers_word',
        'toggle_is_checked',
    ]
    
    def toggle_is_checked(self, request, queryset):
        for record in queryset:
            record.is_checked = not record.is_checked
            record.save()

    toggle_is_checked.short_description = "Toggle is_checked"

    # Download all customer data in excel (xls) format
    def export_selected_customers_xls(self, request, queryset):
        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="customer_records.xls"'

        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('Customer Records')

        row_num = 0

        headers = ['Name', 'Surname', 'Gender', 'Age', 'Mobile Number', 'Malls', 'Brands', 'Created At']
        column_widths = [4000, 4000, 3000, 3000, 5000, 8000, 8000, 5000, 20000]

        for col_num, (header, width) in enumerate(zip(headers, column_widths)):
            worksheet.write(row_num, col_num, header)
            worksheet.col(col_num).width = width

        for customer in queryset:
            row_num += 1
            malls = ', '.join([mall.name for mall in customer.malls.all()])
            brands = ', '.join([brand.name for brand in customer.brands.all()])
            created_at_str = customer.created_at.strftime('%m/%d/%y - %H:%M') 

            data = [
                customer.name,
                customer.surname,
                customer.gender,
                customer.age,
                customer.mobile_number,
                malls,
                brands,
                created_at_str,
            ]

            for col_num, value in enumerate(data):
                worksheet.write(row_num, col_num, value)
                
        workbook.save(response)
        return response

    export_selected_customers_xls.short_description = 'Export selected customers as XLS'

    # Download name_phone customer data in excel (xls) format
    def export_selected_customers_xls_name_phone(self, request, queryset):
        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="customer_name_phone.xls"'

        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('Customer Name and Phone')

        row_num = 0

        headers = ['Name', 'Mobile Number']
        for col_num, header in enumerate(headers):
            worksheet.write(row_num, col_num, header)

        for customer in queryset:
            row_num += 1
            data = [
                customer.name,
                customer.mobile_number,
            ]
            
            for col_num, value in enumerate(data):
                worksheet.write(row_num, col_num, value)

        workbook.save(response)
        return response
                
    export_selected_customers_xls_name_phone.short_description = 'Export selected Customers (Name and Phone) as XLS'
    
    # Download name_phone customer data in csv format
    def export_selected_customers_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="customer_records.csv"'

        writer = csv.writer(response)
        writer.writerow(['Name', 'Mobile Number'])

        for customer in queryset:
            writer.writerow([customer.name, customer.mobile_number])

        return response

    export_selected_customers_csv.short_description = 'Export selected Customers as CSV'

    # Download all customer data in Word format

    def export_selected_customers_word(self, request, queryset):
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="customer_records.docx"'

        document = Document()

        # Iterate over selected customers and add relevant data to the document
        for i, customer in enumerate(queryset):
            if i > 0:
                document.add_paragraph("")  # Add empty paragraph for spacing between customers

            # Add customer heading
            customer_heading = document.add_heading(level=1)
            customer_heading_run = customer_heading.add_run(f'Customer {i+1}: {customer.name}')
            customer_heading_run.bold = True

            # Add customer data
            document.add_paragraph("")  # Add empty paragraph for spacing
            document.add_paragraph(f"Name:    {customer.name}")
            document.add_paragraph(f"Mobile Number: {customer.mobile_number}")
            document.add_paragraph(f"Malls: {', '.join([str(mall) for mall in customer.malls.all()])}")
            document.add_paragraph(f"Brands: {', '.join([brand.name for brand in customer.brands.all()])}")
            document.add_paragraph(f"Created At: {customer.created_at.strftime('%m/%d/%y - %H:%M')}")
            document.add_paragraph("")  # Add empty paragraph for spacing

            # Add Comments section
            document.add_heading('Comments', level=2)
            if customer.comments:
                document.add_paragraph(customer.comments)
            else:
                document.add_paragraph("No comments available.")

        # Save the document to a BytesIO object
        output = io.BytesIO()
        document.save(output)

        # Retrieve the content from the BytesIO object and assign it to the response
        response.content = output.getvalue()
        output.close()

        return response

    export_selected_customers_word.short_description = 'Export selected Customers as Word'

# CustomerRecords, CustomerRecordsAdmin register here.
admin.site.register(CustomerRecords, CustomerRecordsAdmin)

